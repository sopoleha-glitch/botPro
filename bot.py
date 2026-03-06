import asyncio
import logging
import json
import hashlib
import time
import aiosqlite
import aiohttp
import PyPDF2
import io
import pytesseract
from PIL import Image
import pandas as pd
from openpyxl import load_workbook
from datetime import datetime
from typing import Any, Awaitable, Callable, Dict, Optional
from docx import Document
from urllib.parse import quote
from aiogram import Bot, Dispatcher, types, F
from aiogram.enums import ParseMode
from aiogram.client.default import DefaultBotProperties
from aiogram.filters import Command, CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    ReplyKeyboardMarkup, KeyboardButton,
    InlineKeyboardMarkup, InlineKeyboardButton,
    FSInputFile, BufferedInputFile
)
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.filters import StateFilter
from functools import lru_cache, wraps
import asyncio
from concurrent.futures import ThreadPoolExecutor

BOT_TOKEN = "8667653728:AAF3Ekms8refE2-BvS1tgDl03sVuLpvvpx0"
ADMIN_ID = 745613614
DEEPSEEK_API_KEY = "sk-a45c0fa810f4430e8a154955c153070d"
VIP_CHANNEL_URL = "https://t.me/squad_vpotoke"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

DB_PATH = 'tokenbot.db'
executor = ThreadPoolExecutor(max_workers=4)

class DatabaseManager:
    _instance = None
    _connection = None
    _lock = asyncio.Lock()
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    async def get_connection(self):
        if self._connection is None:
            self._connection = await aiosqlite.connect(DB_PATH)
            # Оптимизация SQLite
            await self._connection.execute("PRAGMA journal_mode=WAL")
            await self._connection.execute("PRAGMA synchronous=NORMAL")
            await self._connection.execute("PRAGMA cache_size=10000")
            await self._connection.execute("PRAGMA temp_store=MEMORY")
            await self._connection.execute("PRAGMA mmap_size=30000000000")
        return self._connection
    
    async def close(self):
        if self._connection:
            await self._connection.close()
            self._connection = None
    
    async def execute(self, query: str, params: tuple = ()):
        async with self._lock:
            conn = await self.get_connection()
            try:
                async with conn.execute(query, params) as cursor:
                    await conn.commit()
                    return cursor
            except Exception as e:
                logger.error(f"Database error: {e}")
                raise
    
    async def fetchone(self, query: str, params: tuple = ()):
        async with self._lock:
            conn = await self.get_connection()
            async with conn.execute(query, params) as cursor:
                return await cursor.fetchone()
    
    async def fetchall(self, query: str, params: tuple = ()):
        async with self._lock:
            conn = await self.get_connection()
            async with conn.execute(query, params) as cursor:
                return await cursor.fetchall()
    
    async def executemany(self, query: str, params: list):
        async with self._lock:
            conn = await self.get_connection()
            await conn.executemany(query, params)
            await conn.commit()

db_manager = DatabaseManager()

class CacheManager:
    def __init__(self, ttl: int = 60):
        self.cache = {}
        self.ttl = ttl
        self._cleanup_task = None
    
    async def start_cleanup(self):
        async def cleanup():
            while True:
                await asyncio.sleep(300) 
                now = time.time()
                self.cache = {k: v for k, v in self.cache.items() 
                            if now - v[1] < self.ttl}
        self._cleanup_task = asyncio.create_task(cleanup())
    
    async def get(self, key: str):
        if key in self.cache:
            data, timestamp = self.cache[key]
            if time.time() - timestamp < self.ttl:
                return data
            del self.cache[key]
        return None
    
    async def set(self, key: str, value: Any):
        self.cache[key] = (value, time.time())
    
    async def delete(self, key: str):
        self.cache.pop(key, None)
    
    async def clear(self):
        self.cache.clear()

cache = CacheManager(ttl=300)

def timing_decorator(func):
    @wraps(func)
    async def wrapper(*args, **kwargs):
        start = time.time()
        result = await func(*args, **kwargs)
        duration = time.time() - start
        if duration > 1:
            logger.warning(f"Slow operation {func.__name__}: {duration:.2f}s")
        return result
    return wrapper

async def run_in_executor(func, *args):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, func, *args)

class ScheduleStates(StatesGroup):
    waiting_for_day = State()
    waiting_for_time_task = State()
    waiting_for_week_parity = State()

class ImageStates(StatesGroup):
    waiting_for_prompt = State()

class EditStates(StatesGroup):
    waiting_for_edit_text = State()
    waiting_for_ai_prompt = State()
    waiting_for_excel_cell = State()
    waiting_for_excel_value = State()
    waiting_for_edit_choice = State()
    waiting_for_word_replace = State()

class TokenBotDB:
    
    @staticmethod
    async def init_db():
        await db_manager.execute('''
            CREATE TABLE IF NOT EXISTS users (
                user_id INTEGER PRIMARY KEY,
                username TEXT,
                first_name TEXT,
                tokens INTEGER DEFAULT 10,
                referral_code TEXT UNIQUE,
                referred_by INTEGER,
                total_earned INTEGER DEFAULT 10,
                total_spent INTEGER DEFAULT 0,
                join_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                last_activity TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                is_banned INTEGER DEFAULT 0
            )
        ''')
        
        await db_manager.execute('''
            CREATE TABLE IF NOT EXISTS chat_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                role TEXT,
                content TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        await db_manager.execute('''
            CREATE TABLE IF NOT EXISTS schedule (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                time TEXT,
                day_of_week INTEGER,
                task TEXT,
                week_parity TEXT DEFAULT 'все',
                enabled INTEGER DEFAULT 1,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        await db_manager.execute('''
            CREATE TABLE IF NOT EXISTS schedule_access (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER UNIQUE,
                start_date TEXT,
                active INTEGER DEFAULT 1
            )
        ''')
        
        await db_manager.execute('''
            CREATE TABLE IF NOT EXISTS payments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                amount REAL,
                tokens INTEGER,
                status TEXT DEFAULT 'pending',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                verified_at TIMESTAMP
            )
        ''')
        
        await db_manager.execute('''
            CREATE TABLE IF NOT EXISTS referrals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                referrer_id INTEGER,
                referred_id INTEGER,
                bonus_tokens INTEGER DEFAULT 5,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        await db_manager.execute("CREATE INDEX IF NOT EXISTS idx_chat_history_user ON chat_history(user_id)")
        await db_manager.execute("CREATE INDEX IF NOT EXISTS idx_schedule_user ON schedule(user_id)")
        await db_manager.execute("CREATE INDEX IF NOT EXISTS idx_schedule_day ON schedule(day_of_week, enabled)")
        await db_manager.execute("CREATE INDEX IF NOT EXISTS idx_referrals_referrer ON referrals(referrer_id)")
        await db_manager.execute("CREATE INDEX IF NOT EXISTS idx_users_tokens ON users(tokens)")
        await db_manager.execute("CREATE INDEX IF NOT EXISTS idx_users_referral ON users(referral_code)")
        
        logger.info("Database initialized with optimizations")

    @staticmethod
    @timing_decorator
    async def get_user(user_id: int):
        cache_key = f"user_{user_id}"
        cached = await cache.get(cache_key)
        if cached:
            return cached
        
        result = await db_manager.fetchone(
            "SELECT * FROM users WHERE user_id = ?", 
            (user_id,)
        )
        
        if result:
            await cache.set(cache_key, result)
        return result

    @staticmethod
    @timing_decorator
    async def create_user(user_id: int, username: str, first_name: str, referred_by: int = None):
        referral_code = hashlib.md5(f"{user_id}{time.time()}".encode()).hexdigest()[:8].upper()
        
        conn = await db_manager.get_connection()
        async with conn.execute("BEGIN TRANSACTION"):
            await db_manager.execute(
                "INSERT OR REPLACE INTO users (user_id, username, first_name, referral_code, referred_by) VALUES (?, ?, ?, ?, ?)",
                (user_id, username or "", first_name or "User", referral_code, referred_by)
            )
            if referred_by:
                await db_manager.execute(
                    "UPDATE users SET tokens = tokens + 5 WHERE user_id = ?", 
                    (referred_by,)
                )
                await db_manager.execute(
                    "INSERT INTO referrals (referrer_id, referred_id) VALUES (?, ?)",
                    (referred_by, user_id)
                )
        
        await cache.delete(f"user_{user_id}")
        if referred_by:
            await cache.delete(f"user_{referred_by}")
            await cache.delete(f"referrals_{referred_by}")
        
        return referral_code

    @staticmethod
    @timing_decorator
    async def update_tokens(user_id: int, tokens: int):
        conn = await db_manager.get_connection()
        async with conn.execute("BEGIN TRANSACTION"):
            if tokens > 0:
                await db_manager.execute(
                    "UPDATE users SET tokens = tokens + ?, total_earned = total_earned + ?, last_activity = CURRENT_TIMESTAMP WHERE user_id = ?", 
                    (tokens, tokens, user_id)
                )
            else:
                await db_manager.execute(
                    "UPDATE users SET tokens = tokens + ?, total_spent = total_spent + ?, last_activity = CURRENT_TIMESTAMP WHERE user_id = ?", 
                    (tokens, abs(tokens), user_id)
                )
        
        # Инвалидируем кэш
        await cache.delete(f"user_{user_id}")

    @staticmethod
    @timing_decorator
    async def get_chat_history(user_id: int, limit: int = 20):
        return await db_manager.fetchall(
            "SELECT role, content FROM chat_history WHERE user_id = ? ORDER BY timestamp ASC LIMIT ?",
            (user_id, limit)
        )

    @staticmethod
    @timing_decorator
    async def save_chat_message(user_id: int, role: str, content: str):
        await db_manager.execute(
            "INSERT INTO chat_history (user_id, role, content) VALUES (?, ?, ?)",
            (user_id, role, content)
        )
        
        await db_manager.execute('''
            DELETE FROM chat_history 
            WHERE id IN (
                SELECT id FROM chat_history 
                WHERE user_id = ? 
                ORDER BY timestamp DESC 
                LIMIT -1 OFFSET 100
            )
        ''', (user_id,))

    @staticmethod
    async def get_user_by_referral(code: str):
        return await db_manager.fetchone(
            "SELECT user_id FROM users WHERE referral_code = ?", 
            (code,)
        )

    @staticmethod
    async def get_referrals_count(user_id: int):
        cache_key = f"referrals_{user_id}"
        cached = await cache.get(cache_key)
        if cached is not None:
            return cached
        
        result = await db_manager.fetchone(
            "SELECT COUNT(*) FROM referrals WHERE referrer_id = ?", 
            (user_id,)
        )
        count = result[0] if result else 0
        await cache.set(cache_key, count)
        return count

    @staticmethod
    async def get_schedule_tasks(user_id: int):
        return await db_manager.fetchall(
            "SELECT id, time, day_of_week, task, week_parity, enabled FROM schedule WHERE user_id = ? ORDER BY day_of_week, time",
            (user_id,)
        )

    @staticmethod
    async def add_schedule_task(user_id: int, time: str, day: int, task: str):
        result = await db_manager.execute(
            "INSERT INTO schedule (user_id, time, day_of_week, task) VALUES (?, ?, ?, ?) RETURNING id",
            (user_id, time, day, task)
        )
        row = await result.fetchone()
        await cache.delete(f"schedule_{user_id}")
        return row[0] if row else None

    @staticmethod
    async def delete_schedule_task(task_id: int, user_id: int):
        await db_manager.execute(
            "DELETE FROM schedule WHERE id = ? AND user_id = ?", 
            (task_id, user_id)
        )
        await cache.delete(f"schedule_{user_id}")

    @staticmethod
    async def update_schedule_week_parity(task_id: int, user_id: int, week_parity: str):
        await db_manager.execute(
            "UPDATE schedule SET week_parity = ? WHERE id = ? AND user_id = ?", 
            (week_parity, task_id, user_id)
        )
        await cache.delete(f"schedule_{user_id}")

    @staticmethod
    async def toggle_schedule_task(task_id: int, user_id: int):
        await db_manager.execute(
            "UPDATE schedule SET enabled = NOT enabled WHERE id = ? AND user_id = ?", 
            (task_id, user_id)
        )
        await cache.delete(f"schedule_{user_id}")

    @staticmethod
    async def check_schedule_access(user_id: int):
        result = await db_manager.fetchone(
            "SELECT start_date FROM schedule_access WHERE user_id = ? AND active = 1",
            (user_id,)
        )
        
        if not result:
            await db_manager.execute(
                "INSERT OR REPLACE INTO schedule_access (user_id, start_date, active) VALUES (?, date('now'), 1)",
                (user_id,)
            )
            return True, 0, 0
        
        start_date = datetime.strptime(result[0], '%Y-%m-%d').date()
        days_used = (datetime.now().date() - start_date).days
        if days_used < 0:
            days_used = 0
        
        if days_used <= 13:
            return True, 0, days_used
        else:
            return True, 1, days_used

    @staticmethod
    async def get_today_schedule_users():
        day_of_week = datetime.now().weekday()
        week_parity = get_week_parity()
        
        return await db_manager.fetchall('''
            SELECT DISTINCT user_id FROM schedule 
            WHERE enabled = 1 AND day_of_week = ? 
            AND (week_parity = 'все' OR week_parity = ?)
        ''', (day_of_week, week_parity))

    @staticmethod
    async def cleanup_old_data():
        """Очистка старых данных"""
        await db_manager.execute('''
            DELETE FROM chat_history 
            WHERE timestamp < datetime('now', '-30 days')
        ''')
        
        await db_manager.execute('''
            UPDATE schedule_access 
            SET active = 0 
            WHERE user_id NOT IN (
                SELECT DISTINCT user_id FROM schedule 
                WHERE created_at > datetime('now', '-7 days')
            )
        ''')

_main_keyboard = None

def get_main_keyboard():
    global _main_keyboard
    if _main_keyboard is None:
        _main_keyboard = ReplyKeyboardMarkup(
            keyboard=[
                [KeyboardButton(text="💰 Баланс"), KeyboardButton(text="💳 Купить")],
                [KeyboardButton(text="👥 Рефералы"), KeyboardButton(text="📅 Расписание")],
                [KeyboardButton(text="🎨 Нарисовать"), KeyboardButton(text="📸 OCR фото")],
                [KeyboardButton(text="📄 Редактор файлов"), KeyboardButton(text="📚 Команды")],
                [KeyboardButton(text="🧹 Очистить")]
            ],
            resize_keyboard=True
        )
    return _main_keyboard

def get_schedule_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➕ Добавить", callback_data="schedule_add"),
         InlineKeyboardButton(text="❌ Удалить", callback_data="schedule_del")],
        [InlineKeyboardButton(text="🔄 Чёт/Нечет", callback_data="schedule_week"),
         InlineKeyboardButton(text="🔄 Вкл/Выкл", callback_data="schedule_toggle_list")],
        [InlineKeyboardButton(text="◀️ Назад", callback_data="back_to_main")]
    ])

def get_days_keyboard():
    days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=day, callback_data=f"add_day_{i}") for i, day in enumerate(days)],
        [InlineKeyboardButton(text="◀️ Отмена", callback_data="schedule")]
    ])

def get_file_edit_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✏️ Редактировать вручную", callback_data="edit_manual")],
        [InlineKeyboardButton(text="🤖 AI-помощник", callback_data="edit_ai")],
        [InlineKeyboardButton(text="➕ Добавить в конец", callback_data="edit_append")],
        [InlineKeyboardButton(text="📊 Статистика файла", callback_data="edit_stats")],
        [InlineKeyboardButton(text="📈 Excel: изменить ячейку", callback_data="edit_excel_cell")],
        [InlineKeyboardButton(text="📊 Excel: сортировка", callback_data="edit_excel_sort")],
        [InlineKeyboardButton(text="📝 Word: замена текста", callback_data="edit_word_replace")]
    ])

def get_week_parity():
    week_number = datetime.now().isocalendar()[1]
    return "четная" if week_number % 2 == 0 else "нечетная"

def get_week_parity_russian():
    week_number = datetime.now().isocalendar()[1]
    return "Чётная" if week_number % 2 == 0 else "Нечётная"

def split_long_message(text, max_length=4000):
    if len(text) <= max_length:
        return [text]
    
    parts = []
    while text:
        if len(text) <= max_length:
            parts.append(text)
            break
        
        split_pos = text.rfind('\n', 0, max_length)
        if split_pos == -1:
            split_pos = text.rfind(' ', 0, max_length)
        if split_pos == -1:
            split_pos = max_length
        
        parts.append(text[:split_pos])
        text = text[split_pos:].lstrip()
    
    return parts

def file_size(size_bytes):
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"

@timing_decorator
async def generate_image(prompt: str):
    try:
        encoded_prompt = quote(prompt)
        url = f"https://image.pollinations.ai/prompt/{encoded_prompt}"
        params = {
            "width": 1024,
            "height": 1024,
            "model": "flux",
            "nologo": "true"
        }
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url, params=params, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                if resp.status == 200:
                    return await resp.read()
                else:
                    logger.error(f"Pollinations error: {resp.status}")
                    return None
    except asyncio.TimeoutError:
        logger.error("Timeout generating image")
        return None
    except Exception as e:
        logger.error(f"Image generation error: {e}")
        return None

@timing_decorator
async def ask_deepseek_simple(prompt: str, history=None):
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    messages = []
    if history:
        for role, content in history:
            messages.append({"role": role, "content": content})
    messages.append({"role": "user", "content": prompt})
    
    data = {
        "model": "deepseek-chat",
        "messages": messages,
        "max_tokens": 2000,
        "stream": False,
        "temperature": 0.3
    }
    
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=aiohttp.ClientTimeout(total=30)
            ) as response:
                
                if response.status != 200:
                    error_text = await response.text()
                    logger.error(f"DeepSeek error {response.status}: {error_text}")
                    return f"😔 Ошибка DeepSeek: {response.status}"
                
                result = await response.json()
                full_response = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                
                if not full_response:
                    full_response = "😔 DeepSeek не вернул ответ"
                
                return full_response
                
    except asyncio.TimeoutError:
        logger.error("Таймаут DeepSeek")
        return "⏱️ DeepSeek не отвечает (таймаут 30 сек)"
        
    except Exception as e:
        logger.error(f"DeepSeek error: {e}")
        return f"😔 Ошибка: {str(e)[:50]}"

def sync_ocr(image_bytes):
    """Синхронная функция для OCR"""
    image = Image.open(io.BytesIO(image_bytes))
    return pytesseract.image_to_string(image, lang='rus+eng')

def sync_read_pdf(file_bytes):
    """Синхронная функция для чтения PDF"""
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

async def send_schedule_to_user(bot: Bot, user_id: int):
    try:
        today = datetime.now().strftime("%d.%m.%Y")
        day_of_week = datetime.now().weekday()
        week_parity = get_week_parity()
        days = ["пн", "вт", "ср", "чт", "пт", "сб", "вс"]
        
        has_access, fee, days_used = await TokenBotDB.check_schedule_access(user_id)
        
        if fee > 0:
            user = await TokenBotDB.get_user(user_id)
            if not user or user[3] < 1:
                await bot.send_message(
                    chat_id=user_id,
                    text="❌ Недостаточно токенов для использования расписания. Пополни баланс через /buy",
                    reply_markup=get_main_keyboard()
                )
                return
            await TokenBotDB.update_tokens(user_id, -1)
        
        tasks = await TokenBotDB.get_schedule_tasks(user_id)
        
        today_tasks = []
        for task_id, task_time, task_day, task, task_week_parity, enabled in tasks:
            if task_day == day_of_week and enabled:
                if task_week_parity == "все" or task_week_parity == week_parity:
                    today_tasks.append((task_time, task))
        
        free_days_left = max(0, 14 - days_used) if days_used < 14 else 0
        
        if today_tasks:
            schedule_text = f"🌅 Доброе утро! Твоё расписание на {today} ({days[day_of_week]}, {week_parity} нед.)\n\n"
            for task_time, task in today_tasks:
                schedule_text += f"⏰ {task_time} - {task}\n"
            
            if free_days_left > 0:
                schedule_text += f"\n📅 Бесплатных дней осталось: {free_days_left}"
            else:
                schedule_text += f"\n💰 Списано 1 токен за сегодня"
            
            schedule_text += "\n\nХорошего дня! 🔥"
            
            parts = split_long_message(schedule_text)
            for part in parts:
                await bot.send_message(
                    chat_id=user_id,
                    text=part,
                    reply_markup=get_main_keyboard()
                )
            logger.info(f"Утреннее расписание отправлено пользователю {user_id}")
        
    except Exception as e:
        logger.error(f"Ошибка при отправке расписания пользователю {user_id}: {e}")

async def schedule_checker(bot: Bot):
    last_sent_date = None
    while True:
        try:
            now = datetime.now()
            if now.hour == 9 and now.minute == 0 and last_sent_date != now.date():
                users = await TokenBotDB.get_today_schedule_users()
                
                for i in range(0, len(users), 10):
                    batch = users[i:i+10]
                    tasks = [send_schedule_to_user(bot, user[0]) for user in batch]
                    await asyncio.gather(*tasks, return_exceptions=True)
                    await asyncio.sleep(1) 
                
                last_sent_date = now.date()
            
            if now.hour == 3 and now.minute == 0:
                await TokenBotDB.cleanup_old_data()
                
        except Exception as e:
            logger.error(f"Schedule checker error: {e}")
        
        await asyncio.sleep(60)

class RateLimitMiddleware:
    def __init__(self, rate_limit: int = 1, burst_limit: int = 5):
        self.rate_limit = rate_limit
        self.burst_limit = burst_limit
        self.users = {}
    
    async def __call__(self, handler: Callable, event: types.Message, data: Dict[str, Any]) -> Any:
        user_id = event.from_user.id
        now = time.time()
        
        if user_id in self.users:
            self.users[user_id] = [t for t in self.users[user_id] if now - t < 60]
        else:
            self.users[user_id] = []
        
        if len(self.users[user_id]) >= self.burst_limit:
            await event.answer("⚠️ Слишком много запросов! Подожди минуту.")
            return
        
        if self.users[user_id] and now - self.users[user_id][-1] < self.rate_limit:
            await event.answer("⏳ Подожди секунду между сообщениями!")
            return
        
        self.users[user_id].append(now)
        return await handler(event, data)

bot = Bot(token=BOT_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
dp = Dispatcher(storage=MemoryStorage())

dp.message.middleware(RateLimitMiddleware())

@dp.message(CommandStart())
async def cmd_start(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    username = message.from_user.username or ""
    first_name = message.from_user.first_name or "User"
    
    await state.clear()
    
    args = message.text.split()
    referred_by = None
    if len(args) > 1:
        ref_code = args[1]
        referrer = await TokenBotDB.get_user_by_referral(ref_code)
        if referrer and referrer[0] != user_id:
            referred_by = referrer[0]
    
    user = await TokenBotDB.get_user(user_id)
    if user:
        if user[10]:
            await message.answer("Ты забанен!")
            return
        
        welcome_text = f"""С возвращением, {first_name}!

💰 Твой баланс: {user[3]} токенов
📊 Всего заработано: {user[6]} токенов
💬 Я помню последние 100 сообщений
📁 Могу читать файлы (PDF, Word, TXT, Excel)
📸 Распознаю текст с фото (OCR)
✏️ Редактирую файлы (TXT, DOCX, XLSX)
📅 Можешь создать своё личное расписание
🎨 Генерирую картинки (2 токена)

VIP канал: {VIP_CHANNEL_URL}

Как это работает:
• 1 сообщение = 1 токен
• 1 картинка = 2 токена
• OCR фото = 2 токена
• Редактирование = 1-2 токена
• Загружай файлы - я прочитаю и отвечу
• Приводи друзей (+5 токенов)
• Покупай токены через Stars
• Управляй своим расписанием через 📅 Расписание
• Расписание: 14 дней бесплатно, потом 1 токен/день

Используй кнопки внизу для навигации! 👇"""
    else:
        ref_code = await TokenBotDB.create_user(user_id, username, first_name, referred_by)
        
        welcome_text = f"""Добро пожаловать в TokenBot, {first_name}!

🎁 Бонус: 10 бесплатных токенов!
🔗 Твой реферальный код: {ref_code}
💬 Я помню последние 100 сообщений
📁 Могу читать файлы (PDF, Word, TXT, Excel)
📸 Распознаю текст с фото (OCR)
✏️ Редактирую файлы (TXT, DOCX, XLSX)
📅 Можешь создать своё личное расписание
🎨 Генерирую картинки (2 токена)

Как это работает:
• 10 токенов уже на твоем счету
• 1 сообщение = 1 токен
• 1 картинка = 2 токена
• OCR фото = 2 токена
• Редактирование = 1-2 токена
• Загружай файлы - я прочитаю и отвечу
• Приводи друзей (+5 токенов)
• Управляй своим расписанием через 📅 Расписание
• Расписание: 14 дней бесплатно, потом 1 токен/день

VIP канал: {VIP_CHANNEL_URL}

Используй кнопки внизу для навигации! 👇"""
        
        if referred_by:
            welcome_text += "\n🎊 Твой друг получил 5 токенов!"
    
    await message.answer(welcome_text, reply_markup=get_main_keyboard())

@dp.message(F.text == "💰 Баланс")
@dp.message(Command("balance"))
async def cmd_balance(message: types.Message):
    user = await TokenBotDB.get_user(message.from_user.id)
    if not user:
        await message.answer("Используй /start")
        return
    
    referrals = await TokenBotDB.get_referrals_count(user[0])
    
    bot_username = (await message.bot.me()).username
    text = f"""💰 Твой кошелек

💎 Баланс: {user[3]} токенов
📈 Заработано: {user[6]} токенов
📉 Потрачено: {user[7]} токенов
👥 Рефералов: {referrals}

🔗 Твой код: {user[4]}
📱 Ссылка: https://t.me/{bot_username}?start={user[4]}

Приводи друзей и получай +5 токенов!"""
    
    await message.answer(text, reply_markup=get_main_keyboard())

@dp.message(F.text == "💳 Купить")
@dp.message(Command("buy"))
async def cmd_buy(message: types.Message):
    text = """💳 Магазин токенов

🎯 Предложения:
• 100 токенов - 10 ⭐
• 500 токенов - 45 ⭐ (скидка 10%)
• 1000 токенов - 80 ⭐ (скидка 20%)
• 2000 токенов - 150 ⭐ (скидка 25%)

Оплата через Telegram Stars

Выбери пакет:"""
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="100 - 10⭐", callback_data="buy_100"),
         InlineKeyboardButton(text="500 - 45⭐", callback_data="buy_500")],
        [InlineKeyboardButton(text="1000 - 80⭐", callback_data="buy_1000"),
         InlineKeyboardButton(text="2000 - 150⭐", callback_data="buy_2000")]
    ])
    
    await message.answer(text, reply_markup=keyboard)

@dp.message(F.text == "👥 Рефералы")
@dp.message(Command("referral"))
async def cmd_referral(message: types.Message):
    user = await TokenBotDB.get_user(message.from_user.id)
    if not user:
        await message.answer("Сначала используй /start")
        return
    
    referrals = await TokenBotDB.get_referrals_count(user[0])
    bot_username = (await message.bot.me()).username
    referral_link = f"https://t.me/{bot_username}?start={user[4]}"
    
    text = f"""🔗 Твоя реферальная ссылка

{referral_link}

Статистика:
👥 Приглашено: {referrals}
💰 Заработано: {referrals * 5} токенов

Как это работает:
• Отправляй ссылку друзьям
• Они регистрируются
• Ты получаешь +5 токенов
• Без ограничений!"""
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📱 Поделиться", url=f"https://t.me/share/url?url={referral_link}&text=Заходи в этого бота! Тут можно общаться с ИИ за токены")]
    ])
    
    await message.answer(text, reply_markup=keyboard)

@dp.message(F.text == "🎨 Нарисовать")
@dp.message(Command("draw"))
async def cmd_draw(message: types.Message, state: FSMContext):
    await message.answer("🎨 Отправь описание картинки, например:\nкот в космосе\n\nСтоимость: 2 токена")
    await state.set_state(ImageStates.waiting_for_prompt)

@dp.message(ImageStates.waiting_for_prompt)
async def process_image_prompt(message: types.Message, state: FSMContext):
    prompt = message.text
    
    user = await TokenBotDB.get_user(message.from_user.id)
    if not user or user[3] < 2:
        await message.answer("❌ Недостаточно токенов! Нужно 2 токена")
        await state.clear()
        return
    
    await message.bot.send_chat_action(chat_id=message.chat.id, action="upload_photo")
    await message.answer("🎨 Генерирую картинку... (до 30 секунд)")
    
    image_bytes = await generate_image(prompt)
    
    if image_bytes:
        await message.answer_photo(
            photo=BufferedInputFile(image_bytes, filename="image.jpg"),
            caption=f"🎨 {prompt}"
        )
        await TokenBotDB.update_tokens(message.from_user.id, -2)
        logger.info(f"User {message.from_user.id} generated image: {prompt[:50]}")
    else:
        await message.answer("😔 Не удалось сгенерировать картинку. Попробуй другой запрос.")
    
    await state.clear()

@dp.message(F.text == "📸 OCR фото")
async def ocr_button(message: types.Message):
    await message.answer("📸 Отправь фото с текстом, я распознаю его (стоимость: 2 токена)")

@dp.message(F.photo)
async def handle_photo(message: types.Message):
    user_id = message.from_user.id
    user = await TokenBotDB.get_user(user_id)
    
    if not user or user[3] < 2:
        await message.answer("❌ Недостаточно токенов! Нужно 2 токена")
        return
    
    file = await message.bot.get_file(message.photo[-1].file_id)
    if file.file_size > 5 * 1024 * 1024:  # 5 MB
        await message.answer("❌ Файл слишком большой (макс 5 МБ)")
        return
    
    await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
    await message.answer("🔍 Распознаю текст с фото...")
    
    try:
        file_bytes = await file.download_as_bytearray()
        
        text = await run_in_executor(sync_ocr, file_bytes)
        
        if text.strip():
            parts = split_long_message(f"📝 **Распознанный текст:**\n\n{text}")
            for part in parts:
                await message.answer(part)
            await TokenBotDB.update_tokens(user_id, -2)
        else:
            await message.answer("😔 Не удалось распознать текст. Попробуй другое фото.")
            
    except Exception as e:
        logger.error(f"OCR error: {e}")
        await message.answer("😔 Ошибка при распознавании текста")

@dp.message(F.text == "📄 Редактор файлов")
async def editor_button(message: types.Message):
    await message.answer(
        "📄 **Редактор файлов**\n\n"
        "Отправь мне файл (TXT, DOCX, XLSX), и я помогу:\n"
        "• ✏️ Редактировать вручную\n"
        "• 🤖 Исправить ошибки через AI\n"
        "• 📊 Посмотреть статистику\n"
        "• 📈 Работа с Excel (ячейки, сортировка)\n"
        "• 📝 Работа с Word (замена текста)"
    )

@dp.message(F.document)
async def handle_editable_document(message: types.Message, state: FSMContext):
    file = message.document
    file_name = file.file_name
    file_ext = file_name.split('.')[-1].lower()
    
    allowed_exts = ['txt', 'docx', 'xlsx']
    if file_ext not in allowed_exts:
        await message.answer("❌ Поддерживаются только TXT, DOCX и XLSX файлы")
        return
    
    # Проверка размера файла
    if file.file_size > 10 * 1024 * 1024:  # 10 MB
        await message.answer("❌ Файл слишком большой (макс 10 МБ)")
        return
    
    user_id = message.from_user.id
    user = await TokenBotDB.get_user(user_id)
    if not user or user[3] < 1:
        await message.answer("❌ Недостаточно токенов! Нужен 1 токен")
        return
    
    await message.answer("📥 Скачиваю файл...")
    
    file_obj = await message.bot.get_file(file.file_id)
    file_bytes = await file_obj.download_as_bytearray()
    
    file_info = {
        'filename': file_name,
        'extension': file_ext,
        'content': file_bytes,
        'size': len(file_bytes)
    }
    
    await state.update_data(current_file=file_info)
    
    await message.answer(
        f"✅ Файл '{file_name}' загружен ({file_size(len(file_bytes))})",
        reply_markup=get_file_edit_keyboard()
    )

@dp.callback_query(F.data == "edit_stats")
async def edit_stats(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info:
        await callback.answer("Файл не найден")
        return
    
    ext = file_info['extension']
    content = file_info['content']
    
    if ext == 'txt':
        text = content.decode('utf-8', errors='ignore')
        lines = len(text.split('\n'))
        words = len(text.split())
        chars = len(text)
        
        stats = f"""📊 **Статистика файла**

📄 Имя: {file_info['filename']}
📦 Размер: {file_size(file_info['size'])}
📝 Строк: {lines}
🔤 Слов: {words}
📏 Символов: {chars}"""
        
    elif ext == 'docx':
        doc = await run_in_executor(Document, io.BytesIO(content))
        paragraphs = len(doc.paragraphs)
        text = '\n'.join([p.text for p in doc.paragraphs])
        words = len(text.split())
        
        stats = f"""📊 **Статистика Word-документа**

📄 Имя: {file_info['filename']}
📦 Размер: {file_size(file_info['size'])}
📝 Параграфов: {paragraphs}
🔤 Слов: {words}"""
        
    elif ext == 'xlsx':
        df = await run_in_executor(pd.read_excel, io.BytesIO(content))
        rows, cols = df.shape
        
        stats = f"""📊 **Статистика Excel-файла**

📄 Имя: {file_info['filename']}
📦 Размер: {file_size(file_info['size'])}
📊 Листов: 1
📏 Строк: {rows}
📐 Столбцов: {cols}
🔢 Ячеек: {rows * cols}"""
    
    parts = split_long_message(stats)
    for part in parts:
        await callback.message.edit_text(part)
    await callback.answer()

@dp.callback_query(F.data == "edit_manual")
async def edit_manual_start(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info:
        await callback.answer("Файл не найден")
        return
    
    ext = file_info['extension']
    
    if ext == 'txt':
        content = file_info['content'].decode('utf-8', errors='ignore')
        preview = content[:500] + "..." if len(content) > 500 else content
        await callback.message.edit_text(
            f"📄 **Текущее содержимое:**\n\n{preview}\n\n"
            "✏️ Отправь новый текст для файла"
        )
        await state.set_state(EditStates.waiting_for_edit_text)
        
    elif ext == 'docx':
        doc = await run_in_executor(Document, io.BytesIO(file_info['content']))
        text = '\n'.join([p.text for p in doc.paragraphs])
        preview = text[:500] + "..." if len(text) > 500 else text
        await callback.message.edit_text(
            f"📄 **Текст из документа:**\n\n{preview}\n\n"
            "✏️ Отправь новый текст для документа"
        )
        await state.set_state(EditStates.waiting_for_edit_text)
        
    elif ext == 'xlsx':
        df = await run_in_executor(pd.read_excel, io.BytesIO(file_info['content']))
        await callback.message.edit_text(
            f"📊 **Excel-файл**\n\n"
            f"{df.head(10).to_string()}\n\n"
            "Для редактирования Excel используй кнопку 'Excel: изменить ячейку'"
        )
    
    await callback.answer()

@dp.message(EditStates.waiting_for_edit_text)
async def edit_manual_process(message: types.Message, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    append_mode = data.get('append_mode', False)
    
    if not file_info:
        await message.answer("❌ Файл не найден. Начни заново.")
        await state.clear()
        return
    
    ext = file_info['extension']
    new_content = message.text
    
    if ext == 'txt':
        if append_mode:
            old_content = file_info['content'].decode('utf-8', errors='ignore')
            file_bytes = (old_content + "\n" + new_content).encode('utf-8')
        else:
            file_bytes = new_content.encode('utf-8')
        new_filename = f"edited_{file_info['filename']}"
        
    elif ext == 'docx':
        if append_mode:
            doc = await run_in_executor(Document, io.BytesIO(file_info['content']))
            doc.add_paragraph(new_content)
        else:
            doc = Document()
            for line in new_content.split('\n'):
                doc.add_paragraph(line)
        
        file_bytes = io.BytesIO()
        await run_in_executor(doc.save, file_bytes)
        file_bytes = file_bytes.getvalue()
        new_filename = f"edited_{file_info['filename']}"
    else:
        await message.answer("❌ Неподдерживаемый формат")
        await state.clear()
        return
    
    await message.answer_document(
        document=BufferedInputFile(file_bytes, filename=new_filename),
        caption="✅ Вот твой отредактированный файл!"
    )
    
    await TokenBotDB.update_tokens(message.from_user.id, -1)
    await state.clear()

@dp.callback_query(F.data == "edit_ai")
async def edit_ai_start(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info:
        await callback.answer("Файл не найден")
        return
    
    ext = file_info['extension']
    
    if ext == 'txt':
        content = file_info['content'].decode('utf-8', errors='ignore')
    elif ext == 'docx':
        doc = await run_in_executor(Document, io.BytesIO(file_info['content']))
        content = '\n'.join([p.text for p in doc.paragraphs])
    else:
        await callback.message.edit_text("❌ AI-редактирование пока только для TXT и DOCX")
        return
    
    await state.update_data(edit_content=content, edit_ext=ext)
    
    await callback.message.edit_text(
        "🤖 **AI-редактор**\n\n"
        "Напиши, что нужно сделать с файлом. Например:\n"
        "• 'исправь орфографические ошибки'\n"
        "• 'перепиши в официальном стиле'\n"
        "• 'сделай текст короче'\n"
        "• 'добавь пункты про Python'"
    )
    await state.set_state(EditStates.waiting_for_ai_prompt)

@dp.message(EditStates.waiting_for_ai_prompt)
async def edit_ai_process(message: types.Message, state: FSMContext):
    data = await state.get_data()
    content = data.get('edit_content', '')
    ext = data.get('edit_ext', 'txt')
    prompt = message.text
    
    await message.answer("🤔 Думаю...")
    
    ai_prompt = f"""У меня есть файл с содержимым:
    
{content[:2000]}

Задача: {prompt}

Верни ТОЛЬКО исправленное содержимое файла, без пояснений."""
    
    response = await ask_deepseek_simple(ai_prompt)
    
    if response and "😔" not in response:
        if ext == 'txt':
            file_bytes = response.encode('utf-8')
            await message.answer_document(
                document=BufferedInputFile(file_bytes, filename="ai_edited.txt"),
                caption="✅ Отредактировано с помощью AI!"
            )
        elif ext == 'docx':
            doc = Document()
            for line in response.split('\n'):
                doc.add_paragraph(line)
            file_bytes = io.BytesIO()
            doc.save(file_bytes)
            await message.answer_document(
                document=BufferedInputFile(file_bytes.getvalue(), filename="ai_edited.docx"),
                caption="✅ Отредактировано с помощью AI!"
            )
        
        await TokenBotDB.update_tokens(message.from_user.id, -2)
    else:
        await message.answer("😔 Не удалось обработать запрос")
    
    await state.clear()

@dp.callback_query(F.data == "edit_excel_cell")
async def edit_excel_cell_start(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info or file_info['extension'] != 'xlsx':
        await callback.answer("Это не Excel-файл")
        return
    
    df = await run_in_executor(pd.read_excel, io.BytesIO(file_info['content']))
    await state.update_data(excel_df=df)
    
    await callback.message.edit_text(
        "📈 **Excel: изменение ячейки**\n\n"
        f"{df.head(10).to_string()}\n\n"
        "Введи номер строки и столбца для изменения (например: '1 2' для строки 1, столбца 2)"
    )
    await state.set_state(EditStates.waiting_for_excel_cell)

@dp.message(EditStates.waiting_for_excel_cell)
async def edit_excel_cell_process(message: types.Message, state: FSMContext):
    try:
        row, col = map(int, message.text.split())
        data = await state.get_data()
        df = data.get('excel_df')
        
        await state.update_data(edit_row=row, edit_col=col)
        await message.answer(f"Введи новое значение для ячейки [{row}, {col}]")
        await state.set_state(EditStates.waiting_for_excel_value)
    except:
        await message.answer("❌ Неверный формат. Используй: 'строка столбец' (например: 1 2)")

@dp.message(EditStates.waiting_for_excel_value)
async def edit_excel_value_process(message: types.Message, state: FSMContext):
    data = await state.get_data()
    df = data.get('excel_df')
    row = data.get('edit_row')
    col = data.get('edit_col')
    value = message.text
    
    try:
        df.iat[row-1, col-1] = value
        
        output = io.BytesIO()
        await run_in_executor(df.to_excel, output, index=False)
        output.seek(0)
        
        await message.answer_document(
            document=BufferedInputFile(output.read(), filename="edited.xlsx"),
            caption="✅ Файл обновлён!"
        )
        
        await TokenBotDB.update_tokens(message.from_user.id, -1)
    except Exception as e:
        await message.answer(f"❌ Ошибка: {e}")
    
    await state.clear()

@dp.callback_query(F.data == "edit_excel_sort")
async def edit_excel_sort(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info or file_info['extension'] != 'xlsx':
        await callback.answer("Это не Excel-файл")
        return
    
    df = await run_in_executor(pd.read_excel, io.BytesIO(file_info['content']))
    df_sorted = df.sort_values(by=df.columns[0])
    
    output = io.BytesIO()
    await run_in_executor(df_sorted.to_excel, output, index=False)
    output.seek(0)
    
    await callback.message.answer_document(
        document=BufferedInputFile(output.read(), filename="sorted.xlsx"),
        caption="✅ Файл отсортирован по первому столбцу"
    )
    await TokenBotDB.update_tokens(callback.from_user.id, -1)
    await callback.answer()

@dp.callback_query(F.data == "edit_word_replace")
async def edit_word_replace_start(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info or file_info['extension'] != 'docx':
        await callback.answer("Это не Word-файл")
        return
    
    doc = await run_in_executor(Document, io.BytesIO(file_info['content']))
    text = '\n'.join([p.text for p in doc.paragraphs])
    
    await state.update_data(word_doc=doc, word_text=text)
    preview = text[:500] + "..." if len(text) > 500 else text
    await callback.message.edit_text(
        "📝 **Word: замена текста**\n\n"
        f"Текущий текст:\n{preview}\n\n"
        "Введи что заменить и на что в формате: 'старое|новое'"
    )
    await state.set_state(EditStates.waiting_for_word_replace)

@dp.message(EditStates.waiting_for_word_replace)
async def edit_word_replace_process(message: types.Message, state: FSMContext):
    try:
        old, new = message.text.split('|')
        data = await state.get_data()
        doc = data.get('word_doc')
        
        for paragraph in doc.paragraphs:
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
        
        output = io.BytesIO()
        await run_in_executor(doc.save, output)
        output.seek(0)
        
        await message.answer_document(
            document=BufferedInputFile(output.read(), filename="replaced.docx"),
            caption="✅ Замена выполнена!"
        )
        
        await TokenBotDB.update_tokens(message.from_user.id, -1)
    except Exception as e:
        await message.answer(f"❌ Ошибка: {e}")
    
    await state.clear()

@dp.callback_query(F.data == "edit_append")
async def edit_append(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    file_info = data.get('current_file')
    
    if not file_info:
        await callback.answer("Файл не найден")
        return
    
    await callback.message.edit_text(
        "➕ Отправь текст, который нужно добавить в конец файла"
    )
    await state.set_state(EditStates.waiting_for_edit_text)
    await state.update_data(append_mode=True)
    await callback.answer()

@dp.message(F.text == "📚 Команды")
@dp.message(Command("help"))
async def cmd_help(message: types.Message):
    text = f"""📚 Список всех команд

👤 Основные команды:
/start - Запустить бота и регистрация
/balance - Проверить баланс токенов
/help - Показать это меню

💰 Покупка токенов:
/buy - Открыть магазин токенов

🎨 Генерация картинок:
/draw - Нарисовать картинку (2 токена)

📸 Распознавание текста:
(просто отправь фото с текстом) - 2 токена

📄 Редактирование файлов:
(отправь TXT, DOCX или XLSX файл)

📅 Твоё личное расписание:
/schedule - Управление расписанием
• 14 дней бесплатно
• Потом 1 токен/день
• Чётные/нечётные недели

👥 Реферальная система:
/referral - Твоя реферальная ссылка

VIP канал: {VIP_CHANNEL_URL}

👇 Используй кнопки внизу!"""
    
    await message.answer(text, reply_markup=get_main_keyboard())

@dp.message(F.text == "🧹 Очистить")
@dp.message(Command("clear"))
async def cmd_clear(message: types.Message):
    await db_manager.execute(
        "DELETE FROM chat_history WHERE user_id = ?", 
        (message.from_user.id,)
    )
    
    await message.answer("🧹 История диалога очищена!", reply_markup=get_main_keyboard())

@dp.message(F.text == "📅 Расписание")
@dp.message(Command("schedule"))
async def cmd_schedule(message: types.Message, state: FSMContext):
    await state.clear()
    user_id = message.from_user.id
    
    has_access, fee, days_used = await TokenBotDB.check_schedule_access(user_id)
    
    week_parity = get_week_parity_russian()
    days = ["пн", "вт", "ср", "чт", "пт", "сб", "вс"]
    today = days[datetime.now().weekday()]
    
    all_tasks = await TokenBotDB.get_schedule_tasks(user_id)
    
    day_of_week = datetime.now().weekday()
    current_week_parity = get_week_parity()
    
    today_tasks = []
    for task_id, task_time, task_day, task, task_week_parity, enabled in all_tasks:
        if task_day == day_of_week and enabled:
            if task_week_parity == "все" or task_week_parity == current_week_parity:
                today_tasks.append((task_time, task))
    
    free_days_left = max(0, 14 - days_used) if days_used < 14 else 0
    
    text = f"📅 ТВОЁ РАСПИСАНИЕ\n"
    text += f"└ Сегодня: {today}, {week_parity} неделя\n"
    
    if free_days_left > 0:
        text += f"└ Бесплатно: {free_days_left} дн.\n"
    
    if today_tasks:
        text += "\n🔹 НА СЕГОДНЯ:\n"
        for task_time, task in today_tasks:
            text += f"   ⏰ {task_time} - {task}\n"
    else:
        text += "\n🔹 НА СЕГОДНЯ: задач нет\n"
    
    text += "\n📋 ВСЕ ЗАДАЧИ:\n"
    if all_tasks:
        for task_id, task_time, task_day, task, week_parity_task, enabled in all_tasks:
            status = "✅" if enabled else "❌"
            week_mark = ""
            if week_parity_task == "четная":
                week_mark = " [ч]"
            elif week_parity_task == "нечетная":
                week_mark = " [нч]"
            text += f"{status} ID{task_id}: {days[task_day]} {task_time} - {task}{week_mark}\n"
    else:
        text += "   У тебя пока нет задач\n"
    
    parts = split_long_message(text)
    for part in parts:
        await message.answer(part, reply_markup=get_schedule_keyboard() if part == parts[0] else None)

@dp.callback_query(F.data == "schedule_add")
async def schedule_add_start(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.edit_text(
        "Выбери день недели:",
        reply_markup=get_days_keyboard()
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("add_day_"))
async def schedule_add_day(callback: types.CallbackQuery, state: FSMContext):
    day = int(callback.data.split("_")[2])
    await state.update_data(day=day)
    await state.set_state(ScheduleStates.waiting_for_time_task)
    
    await callback.message.edit_text(
        "Отправь время и задачу в формате:\nЧЧ:ММ Название задачи\n\nНапример:\n09:00 Подъём"
    )
    await callback.answer()

@dp.message(ScheduleStates.waiting_for_time_task)
async def schedule_add_task(message: types.Message, state: FSMContext):
    data = await state.get_data()
    day = data.get('day')
    
    parts = message.text.strip().split(' ', 1)
    if len(parts) != 2:
        await message.answer("❌ Неверный формат. Используй: ЧЧ:ММ Название")
        return
    
    task_time, task = parts
    
    try:
        datetime.strptime(task_time, "%H:%M")
    except:
        await message.answer("❌ Неверный формат времени. Используй ЧЧ:ММ")
        return
    
    task_id = await TokenBotDB.add_schedule_task(message.from_user.id, task_time, day, task)
    await state.clear()
    
    await message.answer(f"✅ Задача добавлена с ID {task_id}!")
    await cmd_schedule(message, state)

@dp.callback_query(F.data == "schedule_del")
async def schedule_del_menu(callback: types.CallbackQuery):
    tasks = await TokenBotDB.get_schedule_tasks(callback.from_user.id)
    
    if not tasks:
        await callback.answer("Нет задач для удаления")
        return
    
    days = ["пн", "вт", "ср", "чт", "пт", "сб", "вс"]
    keyboard = InlineKeyboardBuilder()
    
    for task_id, task_time, task_day, task, week_parity, enabled in tasks:
        week_mark = " [ч]" if week_parity == "четная" else " [нч]" if week_parity == "нечетная" else ""
        btn_text = f"ID{task_id}: {days[task_day]} {task_time} - {task}{week_mark}"
        keyboard.button(text=btn_text, callback_data=f"del_task_{task_id}")
    
    keyboard.button(text="◀️ Назад", callback_data="schedule")
    keyboard.adjust(1)
    
    await callback.message.edit_text(
        "Выбери задачу для удаления:",
        reply_markup=keyboard.as_markup()
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("del_task_"))
async def schedule_del_confirm(callback: types.CallbackQuery):
    task_id = int(callback.data.split("_")[2])
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Да", callback_data=f"del_yes_{task_id}"),
         InlineKeyboardButton(text="❌ Нет", callback_data="schedule_del")]
    ])
    
    await callback.message.edit_text(
        f"Точно удалить задачу ID {task_id}?",
        reply_markup=keyboard
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("del_yes_"))
async def schedule_del_execute(callback: types.CallbackQuery):
    task_id = int(callback.data.split("_")[2])
    await TokenBotDB.delete_schedule_task(task_id, callback.from_user.id)
    
    await callback.answer("✅ Задача удалена")
    await cmd_schedule(callback.message, None)

@dp.callback_query(F.data == "schedule_week")
async def schedule_week_menu(callback: types.CallbackQuery):
    tasks = await TokenBotDB.get_schedule_tasks(callback.from_user.id)
    
    if not tasks:
        await callback.answer("Сначала добавь задачи")
        return
    
    days = ["пн", "вт", "ср", "чт", "пт", "сб", "вс"]
    keyboard = InlineKeyboardBuilder()
    
    for task_id, task_time, task_day, task, week_parity, enabled in tasks:
        week_mark = " [ч]" if week_parity == "четная" else " [нч]" if week_parity == "нечетная" else ""
        btn_text = f"ID{task_id}: {days[task_day]} {task_time} - {task}{week_mark}"
        keyboard.button(text=btn_text, callback_data=f"week_task_{task_id}")
    
    keyboard.button(text="◀️ Назад", callback_data="schedule")
    keyboard.adjust(1)
    
    await callback.message.edit_text(
        "Выбери задачу для настройки чётности/нечётности недели:",
        reply_markup=keyboard.as_markup()
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("week_task_"))
async def schedule_week_set(callback: types.CallbackQuery):
    task_id = int(callback.data.split("_")[2])
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📅 Каждую", callback_data=f"week_set_{task_id}_все"),
         InlineKeyboardButton(text="🔢 Чётную", callback_data=f"week_set_{task_id}_четная")],
        [InlineKeyboardButton(text="🔢 Нечётную", callback_data=f"week_set_{task_id}_нечетная"),
         InlineKeyboardButton(text="◀️ Назад", callback_data="schedule_week")]
    ])
    
    await callback.message.edit_text(
        "На какой неделе выполнять эту задачу?",
        reply_markup=keyboard
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("week_set_"))
async def schedule_week_save(callback: types.CallbackQuery):
    parts = callback.data.split("_")
    task_id = int(parts[2])
    week_parity = parts[3]
    
    await TokenBotDB.update_schedule_week_parity(task_id, callback.from_user.id, week_parity)
    await callback.answer("✅ Настройка сохранена")
    await cmd_schedule(callback.message, None)

@dp.callback_query(F.data == "schedule_toggle_list")
async def schedule_toggle_list(callback: types.CallbackQuery):
    tasks = await TokenBotDB.get_schedule_tasks(callback.from_user.id)
    
    if not tasks:
        await callback.answer("Нет задач")
        return
    
    days = ["пн", "вт", "ср", "чт", "пт", "сб", "вс"]
    keyboard = InlineKeyboardBuilder()
    
    for task_id, task_time, task_day, task, week_parity, enabled in tasks:
        status = "✅" if enabled else "❌"
        week_mark = " [ч]" if week_parity == "четная" else " [нч]" if week_parity == "нечетная" else ""
        btn_text = f"{status} ID{task_id}: {days[task_day]} {task_time} - {task}{week_mark}"
        keyboard.button(text=btn_text, callback_data=f"toggle_task_{task_id}")
    
    keyboard.button(text="◀️ Назад", callback_data="schedule")
    keyboard.adjust(1)
    
    await callback.message.edit_text(
        "Выбери задачу для включения/выключения:",
        reply_markup=keyboard.as_markup()
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("toggle_task_"))
async def schedule_toggle_task(callback: types.CallbackQuery):
    task_id = int(callback.data.split("_")[2])
    await TokenBotDB.toggle_schedule_task(task_id, callback.from_user.id)
    await callback.answer("✅ Статус изменён")
    await schedule_toggle_list(callback)

@dp.callback_query(F.data == "back_to_main")
async def back_to_main(callback: types.CallbackQuery):
    await callback.message.delete()
    await callback.message.answer("Главное меню:", reply_markup=get_main_keyboard())
    await callback.answer()

@dp.callback_query(F.data == "schedule")
async def schedule_callback(callback: types.CallbackQuery):
    await callback.message.delete()
    await cmd_schedule(callback.message, None)

@dp.callback_query(F.data.startswith("buy_"))
async def buy_callback(callback: types.CallbackQuery):
    packages = {
        'buy_100': {'tokens': 100, 'stars': 10},
        'buy_500': {'tokens': 500, 'stars': 45},
        'buy_1000': {'tokens': 1000, 'stars': 80},
        'buy_2000': {'tokens': 2000, 'stars': 150}
    }
    
    if callback.data not in packages:
        await callback.answer("Неверный пакет")
        return
    
    package = packages[callback.data]
    
    prices = [{"label": f"{package['tokens']} Токенов", "amount": package['stars']}]
    
    await callback.bot.send_invoice(
        chat_id=callback.from_user.id,
        title=f"{package['tokens']} Токенов",
        description=f"Покупка {package['tokens']} токенов",
        payload=f"tokens_{package['tokens']}",
        provider_token="",
        currency="XTR",
        prices=prices,
        start_parameter="buy_tokens"
    )
    
    await callback.answer()

@dp.pre_checkout_query()
async def pre_checkout_handler(pre_checkout_q: types.PreCheckoutQuery):
    await pre_checkout_q.answer(ok=True)

@dp.message(F.successful_payment)
async def successful_payment_handler(message: types.Message):
    user_id = message.from_user.id
    payload = message.successful_payment.invoice_payload
    tokens = int(payload.split('_')[1])
    
    await TokenBotDB.update_tokens(user_id, tokens)
    
    await message.answer(
        f"✅ Оплата прошла успешно!\n💰 Начислено: {tokens} токенов",
        reply_markup=get_main_keyboard()
    )

@dp.message(Command("admin"))
async def cmd_admin(message: types.Message):
    if message.from_user.id != ADMIN_ID:
        await message.answer("Только для админа!")
        return
    
    users_count = await db_manager.fetchone("SELECT COUNT(*) FROM users")
    users = users_count[0] if users_count else 0
    
    tokens_sum = await db_manager.fetchone("SELECT SUM(tokens) FROM users")
    tokens = tokens_sum[0] or 0 if tokens_sum else 0
    
    schedule_count_res = await db_manager.fetchone("SELECT COUNT(*) FROM schedule")
    schedule_count = schedule_count_res[0] if schedule_count_res else 0
    
    pending_payments_res = await db_manager.fetchone("SELECT COUNT(*) FROM payments WHERE status='pending'")
    pending_payments = pending_payments_res[0] if pending_payments_res else 0
    
    banned_users_res = await db_manager.fetchone("SELECT COUNT(*) FROM users WHERE is_banned=1")
    banned_users = banned_users_res[0] if banned_users_res else 0
    
    text = f"""📊 АДМИН ПАНЕЛЬ

👥 Пользователи: {users}
💰 Всего токенов: {tokens}
📅 Задач в расписании: {schedule_count}
⏳ Ожидают платежей: {pending_payments}
🚫 Забанено: {banned_users}

Обновлено: {datetime.now().strftime('%d/%m/%Y %H:%M')}"""
    
    await message.answer(text)

async def shutdown():
    """Graceful shutdown"""
    logger.info("Shutting down...")
    await db_manager.close()
    executor.shutdown(wait=True)
    await cache.clear()

async def main():
    try:
        await TokenBotDB.init_db()
        await cache.start_cleanup()
        
        asyncio.create_task(schedule_checker(bot))
        
        logger.info("Бот запущен с оптимизированной БД!")
        await dp.start_polling(bot)
    finally:
        await shutdown()

if __name__ == "__main__":
    asyncio.run(main())
